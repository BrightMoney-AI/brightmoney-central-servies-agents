"""
generate_design_doc.py — Full engineering design document for the L0 Metrics Report system.
Run: python generate_design_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "L0_Metrics_System_Design.docx"

# ── Palette ────────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1A, 0x56, 0xDB)
DARK       = RGBColor(0x1F, 0x25, 0x37)
MID        = RGBColor(0x37, 0x41, 0x51)
LIGHT      = RGBColor(0x6B, 0x72, 0x80)
MONO_BG    = RGBColor(0xF3, 0xF4, 0xF6)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
ROW_EVEN   = "EFF6FF"
ROW_ODD    = "FFFFFF"
HDR_COLOR  = "1A56DB"
CODE_BG    = "F3F4F6"

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


# ── XML helpers ────────────────────────────────────────────────────────────────
def _shd(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


# ── Style functions ────────────────────────────────────────────────────────────
def h1(text):
    p   = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = BLUE
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)


def h2(text):
    p   = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = DARK
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def h3(text):
    p   = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = MID
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)


def body(text, bold=False, italic=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(10.5)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)


def bullet(text, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)


def code(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x25, 0x37)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)


def table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _shd(hdr_cells[i], HDR_COLOR)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold      = True
                run.font.color.rgb = WHITE
                run.font.size      = Pt(9.5)
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg    = ROW_EVEN if ri % 2 == 0 else ROW_ODD
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            _shd(cells[ci], bg)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    if col_widths:
        for r in tbl.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()


def divider():
    p   = doc.add_paragraph()
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xDB, 0xEA, 0xFE)
    run.font.size      = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(50)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("L0 Metrics Report System")
run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Comprehensive Engineering Design Document")
run2.font.size = Pt(14); run2.font.italic = True; run2.font.color.rgb = MID

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Bright Money  ·  Engineering Platform  ·  Version 1.0  ·  2026")
run3.font.size = Pt(11); run3.font.color.rgb = LIGHT

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. System Overview")
body(
    "The L0 Metrics Report is an automated observability pipeline that runs daily at 10:00 IST. "
    "It collects infrastructure health, API performance, and business metrics for all Bright Money "
    "engineering teams, then publishes structured Slack Canvases before the morning standup — "
    "reducing review time from 15–20 minutes to under 2 minutes."
)
body("The system produces six distinct Slack Canvases per run:")

table(
    ["Canvas Title",                        "Group",            "Content"],
    [
        ["{Group} — L0 Daily Metrics",       "All groups",       "Per-service system health, API metrics, per-endpoint breakdown, queue depths"],
        ["Central Services — Business Metrics","Central Services","PromQL-backed business event counters from central_business.json"],
        ["UAA Services — Business Metrics",  "UAA Services",     "Trino/Iceberg: Onboarding, Account Linking, Plaid, Partner Costs, SAISM/ALSM latency"],
        ["Data Platform — Business Metrics", "Data Platform",    "Trino: CDC recency, compaction, view staleness, Debezium validation"],
        ["Data Platform — EMR Metrics",      "Data Platform",    "Trino: Cube health, staleness, memory, CPU, schedule delay, row growth, execution time"],
        ["Data Platform — connector health", "Data Platform",    "VictoriaMetrics: CDC sink lag, coordinator status, Kafka sink metrics, VM disk"],
    ],
    col_widths=[2.4, 1.6, 3.8],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 2. REPOSITORY STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Repository Structure")

table(
    ["File / Directory",                     "Responsibility"],
    [
        ["main.py",                           "CLI entry point — --now flag triggers immediate run; default starts the local APScheduler"],
        ["scheduler.py",                      "Core orchestrator — runs all collectors, assembles reports, posts all canvases; defines Block Kit summary builders"],
        ["config.py",                         "Pydantic Settings — all environment variables resolved here; vm_base_url and vm_headers are computed properties"],
        ["services.py",                       "ServiceDef dataclass + loader — reads services.json and ems.json; builds PromQL label selectors per service"],
        ["collector.py",                      "MetricsReport collector — runs all PromQL queries for one service through the gateway; produces MetricsReport"],
        ["formatter.py",                      "to_l0_report() — converts raw MetricsReport into typed L0Report; applies baseline comparison logic"],
        ["spike_analyzer.py",                 "Spike detection — consecutive 2× jump and 1.5× max/avg ratio over 48 × 30-min buckets"],
        ["gateway.py",                        "MetricsGateway — serialises VM queries (asyncio.Lock); 5s timeout per query; non-raising failure collection"],
        ["vm_client.py",                      "VMClient — async httpx wrapper for VictoriaMetrics; query(), query_vector(), query_range()"],
        ["trino_client.py",                   "Trino client — async wrapper over blocking trino.dbapi; thread executor; retry on queue-full and connection errors"],
        ["airflow_client.py",                 "Airflow DB + REST API client — DAG run states via PostgreSQL; view flow health via REST API"],
        ["kafka_connect.py",                  "Kafka Connect health — polls /connectors and /connectors/{n}/status per instance"],
        ["models.py",                         "All dataclasses — Status, Server, Endpoint, L0Report, AirflowDagRun, KafkaConnect*, ViewFlow*"],
        ["canvas_renderer.py",                "render_canvas() — renders list[L0Report] as Slack Canvas markdown; Airflow/Connector/Queue sections"],
        ["renderer.py",                       "render() — original Block Kit renderer (kept for backward compat); shared helpers used by canvas_renderer"],
        ["slack_publisher.py",                "publish_canvas() — creates Slack canvas via canvases.create API; posts summary + canvas URL"],
        ["central_business_collector.py",     "Business metrics from central_business.json — runs PromQL queries, returns list[BusinessMetric]"],
        ["central_business_renderer.py",      "Renders Central Services business metrics canvas"],
        ["uaa_business_collector.py",         "UAA Trino business metrics — Onboarding, Plaid, Account Linking, SAISM/ALSM latency, Partner Costs"],
        ["uaa_business_renderer.py",          "Renders UAA business metrics canvas"],
        ["dp_business_collector.py",          "DP Trino business metrics — CDC recency, compaction, view staleness, Debezium validation"],
        ["dp_business_renderer.py",           "Renders Data Platform business metrics canvas"],
        ["emr_collector.py",                  "EMR/Cosmos cube metrics — 8 concurrent Trino queries; EmrReport with EmrSection per query"],
        ["emr_renderer.py",                   "Renders EMR metrics canvas"],
        ["dp_l0_collector.py",                "DP connector health — CDC sink lag, coordinator, Kafka sinks, VM disk via VictoriaMetrics"],
        ["dp_l0_renderer.py",                 "Renders DP connector health canvas"],
        ["queries/__init__.py",               "PromQL builder functions (build_system_queries, build_api_queries, etc.) + SQL file loaders (load_uaa, load_dp)"],
        ["queries/uaa/*.sql",                 "12 UAA SQL queries — onboarding, account linking, Plaid batch/force refresh, partner costs, txn quality"],
        ["queries/dp/*.sql",                  "17 DP SQL queries — EMR cube health, staleness, memory, CPU, row growth, CDC validation, view recency"],
        ["services.json",                     "30 service definitions — name patterns, job labels, excluded endpoints, RabbitMQ queues, report groups"],
        ["central_business.json",             "Central Services business metric queries — PromQL + display metadata + per-metric thresholds"],
        ["buildspec.yaml",                    "AWS CodeBuild spec — builds Docker image, pushes to ECR on every git push to main"],
    ],
    col_widths=[2.5, 5.3],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 3. CONFIGURATION (config.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Configuration — config.py")

body(
    "All runtime configuration is managed by a single Pydantic Settings class. "
    "Variables are read from environment (or .env file in development). "
    "vm_base_url and vm_headers are computed @property fields — not directly settable."
)

table(
    ["Environment Variable",       "Type",   "Default / Value",                             "Purpose"],
    [
        ["VM_INSTANCE_ENTRYPOINT", "str",    "http://vmselect-observability.brightmoney.net:8481", "VictoriaMetrics cluster base URL"],
        ["VM_INSTANCE_TYPE",       "str",    "cluster",                                      "cluster → appends /select/0/prometheus; single → bare URL"],
        ["VM_AUTH_HEADER",         "str",    "Basic bWNwOlh6UjdY…",                          "Value for the Authorization HTTP header sent to VM"],
        ["GATEWAY_TIMEOUT_SECS",   "float",  "5.0",                                          "Per-query timeout in MetricsGateway"],
        ["QUERY_WINDOW",           "str",    "24h",                                          "PromQL window for all L0 queries"],
        ["SLACK_BOT_TOKEN",        "str",    "(required)",                                   "Bot OAuth token for Slack API calls"],
        ["SLACK_CHANNEL_ID",       "str",    "(required)",                                   "Target Slack channel for all canvas posts"],
        ["TRINO_HOST",             "str",    "int-trino.brightmoney.co",                     "Trino coordinator hostname"],
        ["TRINO_PORT",             "int",    "443",                                          "Trino HTTPS port"],
        ["TRINO_USER",             "str",    "uaa_team_metrics",                             "Trino user for Iceberg queries"],
        ["TRINO_SOURCE",           "str",    "engg_team_code",                               "Trino source tag for query attribution"],
        ["AIRFLOW_DB_URL",         "str",    "(empty = disabled)",                           "PostgreSQL DSN for Airflow metadata DB direct queries"],
        ["AIRFLOW_API_URL",        "str",    "(empty = disabled)",                           "Airflow REST API base URL for view flow DAG"],
        ["AIRFLOW_API_USERNAME",   "str",    "(empty = disabled)",                           "Airflow REST API basic auth username"],
        ["AIRFLOW_API_PASSWORD",   "str",    "(empty = disabled)",                           "Airflow REST API basic auth password"],
        ["KAFKA_CONNECT_KAFKA_SINK_URL", "str", "(empty = disabled)",                        "Kafka Sink connect REST API base URL"],
        ["KAFKA_CONNECT_CDC_SINK_URL",   "str", "(empty = disabled)",                        "CDC Sink connect REST API base URL"],
        ["KAFKA_CONNECT_DEBEZIUM_URL",   "str", "(empty = disabled)",                        "Debezium connect REST API base URL"],
    ],
    col_widths=[2.2, 0.6, 2.2, 2.8],
)

body("Computed properties:")
bullet("vm_base_url: if VM_INSTANCE_TYPE == 'cluster' → {entrypoint}/select/0/prometheus, else bare entrypoint")
bullet("vm_headers: {'Authorization': VM_AUTH_HEADER} if set, else {}")
bullet("kafka_connect_instances: dict of {display_name: url} for non-empty Kafka Connect URLs")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 4. SERVICE CONFIGURATION (services.json)
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Service Configuration — services.json + ems.json")

body(
    "Every monitored service is described by a ServiceDef object. The 30 entries in services.json "
    "define how each service's PromQL selectors are constructed and which report group (canvas) it belongs to."
)

table(
    ["Field",                  "Type",       "Purpose"],
    [
        ["display_name",       "str",        "Human-readable service name shown in the canvas heading"],
        ["report_group",       "str",        "Determines which canvas: 'UAA Services', 'Central Services', 'Data Platform'"],
        ["name_patterns",      "list[str]",  "Regex patterns joined as name=~\"p1|p2\" to scope node_exporter queries"],
        ["system_job",         "str",        "job= label for Prometheus node_exporter scrapes (usually 'system_metrics')"],
        ["api_job",            "str",        "job= label for Django statsd scrapes (usually 'platform_statsd_metrics')"],
        ["api_name_patterns",  "list[str]",  "Overrides name_patterns for API/endpoint queries when naming differs"],
        ["api_exclude_endpoints","list[str]","Endpoints excluded from per-endpoint breakdown (noise filtering)"],
        ["api_method",         "str",        "method= label filter; e.g. 'POST' for POST-only services"],
        ["api_request_metric", "str",        "Override for request counter (default: django_request_count)"],
        ["api_response_metric","str",        "Override for response counter (default: django_http_responses_total_by_status)"],
        ["rabbitmq_queues",    "list[str]",  "Queue names to pull from RabbitMQ Prometheus exporter"],
    ],
    col_widths=[2.0, 1.2, 4.6],
)

body(
    "ems.json is a Grafana dashboard export. parse_ems_dashboard() extracts one ServiceDef per "
    "row section by scanning all panel targets for name=~\"...\" patterns. Services from both files "
    "are merged by display_name — later entries override jobs, name_patterns are unioned."
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 5. EXECUTION FLOW (scheduler.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Execution Flow — scheduler.py")

body(
    "run_report() is the single async function that orchestrates everything. "
    "It is triggered by the Airflow KubernetesPodOperator at 10:00 IST (04:30 UTC)."
)

h2("5.1  Phase 1 — VM Collection (inside async with VMClient)")
body("A single VMClient context is opened for the entire VM collection phase.")

bullet("Step 1: load_services() reads services.json + ems.json, merges, optionally filters by --group")
bullet("Step 2: MetricsGateway is instantiated with timeout_secs=5.0")
bullet("Step 3: For each service, collector.collect(vm, gateway, service) runs all PromQL queries serially through the gateway")
bullet("Step 4: formatter.to_l0_report() converts the raw MetricsReport into a typed L0Report")
bullet("Step 5: The L0Report is appended to groups[service.report_group]")
bullet("Step 6: dp_l0_collector.collect_dp_l0() runs CDC/Kafka sink VM queries (inside same VM context)")
bullet("Step 7: central_business_collector.collect_business_metrics(vm) runs PromQL business queries")

h2("5.2  Phase 2 — Trino Collection (outside VMClient)")
body("Trino queries run in a separate thread pool — they are blocking I/O and independent of VM.")

bullet("uaa_business_collector.collect_uaa_business_metrics() — 12 SQL queries concurrently via asyncio.gather")
bullet("dp_business_collector.collect_dp_business_metrics() — 8 SQL queries concurrently")
bullet("emr_collector.collect_emr_metrics() — 8 EMR SQL queries concurrently")

h2("5.3  Phase 3 — Auxiliary Data (concurrent)")
body("Three sources are fetched simultaneously via asyncio.gather:")

bullet("fetch_all_connector_health() — Kafka Connect REST API for all configured instances")
bullet("fetch_airflow_health() — PostgreSQL DAG run states (direct DB query)")
bullet("fetch_view_flow_health() — Airflow REST API for dp_cosmos_execute_view_flow last 24h")

h2("5.4  Phase 4 — Canvas Rendering + Publishing")
body("Canvases are published in a fixed order: UAA Services → Central Services → Data Platform.")

bullet("For each group: render_canvas() → publish_canvas() (summary blocks + canvas URL)")
bullet("Business metrics canvases follow in order: Central → UAA → DP → EMR → DP connector health")
bullet("Each publish_canvas() call: (1) canvases.create API, (2) auth_test for workspace URL, (3) chat_postMessage summary, (4) chat_postMessage canvas URL for unfurl")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 6. DATA COLLECTION LAYER
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Data Collection Layer")

h2("6.1  VMClient (vm_client.py)")
body(
    "Thin async wrapper over VictoriaMetrics' Prometheus-compatible HTTP API. "
    "Uses httpx.AsyncClient as the underlying transport. Must be used as an async context manager."
)

table(
    ["Method",          "API Endpoint",          "Returns",                         "Use Case"],
    [
        ["query()",      "/api/v1/query",          "Optional[float] — scalar",       "Any single-value aggregate: throughput, success rate, latency, baselines"],
        ["query_vector()","/api/v1/query",         "list[(label, float)] — vector",  "Per-server or per-endpoint breakdown; id_label='name' or 'endpoint'"],
        ["query_range()", "/api/v1/query_range",   "list[float] — 48 buckets",       "Spike analysis: 30m step over 24h, NaN/Inf filtered out"],
    ],
    col_widths=[1.5, 1.9, 2.2, 2.2],
)

body(
    "query_range() computes start = now − 24h, end = now, step = 30m. "
    "When combined with rate([30m]) in the PromQL, each bucket is a non-overlapping 30-minute window — "
    "essential for the spike analyzer to fire correctly."
)

h2("6.2  MetricsGateway (gateway.py)")
body(
    "The gateway sits between the collector and VMClient. It solves two production concerns:"
)
bullet("asyncio.Lock serialisation — only one VM query in flight at a time, preventing cluster overload")
bullet("asyncio.wait_for(timeout=5s) — a hung query never blocks the entire collection run")
bullet("Non-raising failure collection — on any error, a FailedQuery is recorded and None is returned")
body(
    "The collector treats None as 'no data' — metrics show as N/A in the canvas rather than crashing. "
    "All failed queries are logged and reported, but the rest of the report is unaffected."
)

h2("6.3  PromQL Query Builders (queries/__init__.py)")
body("Five builder functions generate all PromQL dynamically, scoped to a service's selectors:")

table(
    ["Builder Function",          "Queries Built",                              "Notes"],
    [
        ["build_system_queries()", "cpu_usage_pct, memory_usage_pct, disk_usage_pct, servers_up, servers_down", "per_server=True for cpu/mem/disk; False for up/down counts"],
        ["build_api_queries()",   "throughput_rps, success_rate_pct, error_rate_pct, avg_latency_ms + 3 baselines", "7d baselines use offset 24h to end yesterday"],
        ["build_per_endpoint_queries()", "endpoint_hits, success_pct, error_count, p99_latency_ms + 2 baselines", "by (endpoint) grouping; id_label='endpoint' for query_vector"],
        ["build_queue_queries()", "queue_ready, queue_unacked, queue_total",   "One query per metric; queue names joined as regex"],
        ["build_spike_queries()", "cpu, memory, error_rate, throughput, latency (range)", "step=30m; used with query_range() in collector"],
    ],
    col_widths=[1.9, 3.1, 2.8],
)

h2("6.4  Trino Client (trino_client.py)")
body(
    "Wraps the blocking trino.dbapi driver for async use. "
    "execute_query() offloads the blocking I/O via run_in_executor(None, ...)."
)
bullet("Connection: HTTPS, catalog=iceberg, schema=iceberg_db, request_timeout=300s")
bullet("Retry on TrinoQueryError 'Too many queued queries': up to 3 attempts, linear backoff (15s, 30s, 45s)")
bullet("Retry on TrinoConnectionError: 3 attempts, 15s backoff")
bullet("Connection explicitly closed in finally block after every query — no connection leaks")

h2("6.5  Collector (collector.py)")
body(
    "collect(vm, gateway, service) runs all queries for one service and returns a MetricsReport. "
    "It distinguishes three query modes:"
)
bullet("System queries (per_server=True) → gateway.fetch() → vm.query_vector() → stored in server_values")
bullet("API queries (per_server=False) → gateway.fetch() → vm.query() → stored in values")
bullet("Endpoint queries (per_server=True, id_label='endpoint') → stored in endpoint_values (only when api_job is set)")
bullet("Queue queries → stored in queue_values (only when rabbitmq_queues is configured)")
bullet("Spike queries → vm.query_range() directly (failures are warnings only, do not affect report)")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 7. REPORT ASSEMBLY (formatter.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Report Assembly — formatter.py")

body(
    "to_l0_report() converts a raw MetricsReport into a fully typed L0Report. "
    "This is the central analysis step — it applies all status logic, baseline comparisons, "
    "and produces the final structure that the renderer consumes."
)

h2("7.1  Server Status")
table(
    ["Metric",    "Warning",  "Critical"],
    [
        ["CPU",    "≥ 70%",   "≥ 90%"],
        ["Memory", "≥ 75%",   "≥ 90%"],
        ["Disk",   "≥ 80%",   "≥ 90%"],
    ],
    col_widths=[1.5, 1.5, 1.5],
)

h2("7.2  API Status — Baseline-Relative Comparison")
body(
    "All API metrics are compared against their 7-day baseline (same metric over 7d ending 24h ago). "
    "This catches regressions that absolute thresholds alone would miss."
)

table(
    ["Function",                 "Baseline condition",          "Warning",                         "Critical"],
    [
        ["_latency_spike_icon()", "baseline available",          "current / baseline ≥ 1.5×",      "current / baseline ≥ 2.0×"],
        ["_latency_spike_icon()", "no baseline",                 "current ≥ 500ms",                 "current ≥ 1000ms"],
        ["_success_drop_icon()", "baseline available",           "drop from baseline ≥ 5pp",       "drop from baseline ≥ 10pp"],
        ["_success_drop_icon()", "no baseline",                  "< 95%",                           "< 90%"],
        ["_error_spike_icon()",  "baseline ≥ 0.5%",              "current / baseline ≥ 2×",        "current / baseline ≥ 3×"],
        ["_error_spike_icon()",  "baseline near zero",           "≥ 1% absolute",                  "≥ 5% absolute"],
    ],
    col_widths=[1.9, 1.8, 2.0, 2.1],
)

h2("7.3  Per-Endpoint Status")
body(
    "Only endpoints with ≥ 100 hits are included in the service status computation "
    "(to avoid statistical noise from low-traffic paths). "
    "The same baseline comparison functions apply per endpoint using their individual 7-day baselines."
)

h2("7.4  Overall Service Status")
body(
    "The worst icon across server_icons + api_icons + ep_icons determines the service status. "
    "If show_api_metrics is False (no api_job configured), only server_icons are used. "
    "All icons = ⚪ (all queries returned None) → UNKNOWN status."
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 8. SPIKE DETECTION (spike_analyzer.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Spike Detection — spike_analyzer.py")

body(
    "The spike analyzer receives a list of 48 × 30-minute bucket values and applies two independent "
    "detection algorithms. 24-hour averages alone smooth out transient spikes — this module catches them."
)

h2("8.1  Consecutive Spike (is_spiked)")
body(
    "For each pair of adjacent non-zero buckets, if bucket[i] / bucket[i-1] ≥ 2.0, a spike is recorded. "
    "spike_count tracks how many such jumps occurred; worst_jump tracks the highest ratio seen."
)
body("Rules:")
bullet("Both buckets must be non-zero — a 0 → X transition is not a spike (service starting up)")
bullet("Zero buckets are skipped in the iteration to avoid false positives during quiet periods")

h2("8.2  Elevated Max/Avg Ratio (is_elevated)")
body(
    "max(non_zero) / avg(non_zero) ≥ 1.5 flags the metric as elevated — even without a sharp 2× jump. "
    "This catches gradual elevation: CPU at 40% most of the day then 70% for 2 hours = ratio 1.75×."
)
body("Only fires when is_spiked is False — the two conditions are mutually exclusive in the report label.")

h2("8.3  Output")
table(
    ["SpikeResult Field", "Meaning"],
    [
        ["is_spiked",     "True when at least one consecutive 2× jump was found"],
        ["is_elevated",   "True when max/avg ≥ 1.5 and no spike"],
        ["spike_count",   "Number of 2× consecutive jumps"],
        ["worst_jump",    "Highest bucket[i]/bucket[i-1] ratio seen"],
        ["max_avg_ratio", "max(non_zero) / avg(non_zero) — catches sustained elevation"],
        ["max_val / avg_val / unit", "Used for fmt_max() / fmt_avg() display with correct unit suffix"],
    ],
    col_widths=[1.8, 6.0],
)

body("Minimum requirement: at least 2 non-zero buckets; otherwise analyze() returns None.")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 9. AIRFLOW INTEGRATION (airflow_client.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Airflow Integration — airflow_client.py")

body(
    "Two separate mechanisms fetch Airflow health — a direct PostgreSQL query and the REST API."
)

h2("9.1  PostgreSQL Direct Query (_fetch_sync)")
body(
    "SQLAlchemy with pool_size=1 opens a connection to the Airflow metadata DB. "
    "Two queries run in a single connection:"
)

h3("_QUERY — Latest run per dp_* DAG")
body("Fetches the most recent dag_run row for dp_cosmos_flag_debezium_invalid_tables using DISTINCT ON (dag_id) ORDER BY start_date DESC.")

h3("_PIPELINE_QUERY — Today + Yesterday per pipeline DAG")
body("Fetches the latest run for two pipeline DAGs per IST calendar date:")
bullet("transaction_parallel_ingestion_new_emr")
bullet("cost_cube_pipeline_new_emr")
body(
    "Uses PostgreSQL DISTINCT ON (dag_id, IST_date) to get the most recent attempt per day. "
    "IST bucketing: ((start_date AT TIME ZONE 'UTC') AT TIME ZONE 'Asia/Kolkata')::date. "
    "Window: NOW() - INTERVAL '2 days' to cover today + yesterday in any timezone."
)
body("The run_date field on AirflowDagRun carries the IST calendar date for today/yesterday comparison in the canvas.")

h2("9.2  REST API — View Flow Health")
body(
    "fetch_view_flow_health() calls /api/v1/dags/dp_cosmos_execute_view_flow/dagRuns "
    "with limit=500, start_date_gte=now-24h. "
    "Runs are classified as success / failed / running. "
    "Failed runs surface their table_name from the dag_run conf payload "
    "(base_table_name → table_name → dataset_name → run_id as fallback)."
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 10. KAFKA CONNECT INTEGRATION (kafka_connect.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Kafka Connect Integration — kafka_connect.py")

body(
    "Three Kafka Connect instances are polled (Kafka Sink, CDC Sink, Debezium) via their REST APIs. "
    "All instances are queried concurrently via asyncio.gather."
)
body("Per instance:")
bullet("GET /connectors — fetch list of all connector names")
bullet("GET /connectors/{name}/status — fetch connector state and per-task states")
bullet("A connector is 'unhealthy' if state != RUNNING or any task state != RUNNING")
bullet("On connection failure, the entire instance is skipped (logged as error, not crash)")
body("Result: KafkaConnectHealth → list[KafkaConnectInstance(name, total, unhealthy)]")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 11. CANVAS RENDERING (canvas_renderer.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Canvas Rendering — canvas_renderer.py")

body(
    "render_canvas() produces a single markdown string published to Slack as a Canvas. "
    "The document structure is fixed: service sections → Queue Metrics → Connector Health → Airflow DAGs."
)

h2("11.1  Per-Service Section (_render_service)")
table(
    ["Sub-section",              "Content"],
    [
        ["## Service — 🟢 HEALTHY", "IST timestamp (day + time) in italic below heading"],
        ["### System Health",       "group summary (avg CPU/MEM/Disk); per-server bullet list"],
        ["### API Metrics",         "Throughput · Success rate · Error rate · P50 latency with % vs 7d baseline"],
        ["### API Endpoints · N",   "⚠️ Flagged first (reason label); Top N unflagged by hits; +N more hidden"],
    ],
    col_widths=[2.2, 5.6],
)

h2("11.2  Endpoint Flagging Logic")
body("An endpoint is shown under the ⚠️ Flagged section if any of these hold:")

table(
    ["Condition",                     "Threshold"],
    [
        ["P99 vs 7d baseline",         "≥ 1.5× → Warning; ≥ 2.0× → Critical"],
        ["P99 absolute",               "≥ 1000ms → Warning; ≥ 3000ms → Critical"],
        ["Success rate drop vs baseline","≥ 5pp → Warning; ≥ 10pp → Critical"],
        ["Success rate absolute",       "< 95% → Warning; < 80% → Critical"],
    ],
    col_widths=[2.5, 5.3],
)

h2("11.3  Queue Metrics Section (_render_queue_section)")
body("One ### per service with RabbitMQ queues. Flags: ready ≥ 500 → ⚠️; ready ≥ 100 → 🟡.")

h2("11.4  Connector Health Section (_render_connector_section)")
body("One ### per Kafka Connect instance. Only unhealthy connectors shown in table; healthy → '✅ All N healthy'.")

h2("11.5  Airflow DAGs Section (_render_airflow_section)")
body("Three sub-sections:")
bullet("Pipeline DAGs: today vs yesterday table — state emoji + start time from run_index[(dag_id, IST_date)]")
bullet("dp_* DAG runs: latest state table from _QUERY")
bullet("dp_cosmos_execute_view_flow: 24h summary + failed table refreshes with start times")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 12. SLACK PUBLISHING (slack_publisher.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Slack Publishing — slack_publisher.py")

body("publish_canvas() follows a 4-step flow:")
bullet("Step 1: canvases.create — posts the full markdown as a Slack Canvas document via api_call(json=...) to avoid HTTP 414 on large documents")
bullet("Step 2: auth_test — resolves workspace URL to build deep-link: {workspace}/docs/{team_id}/{canvas_id}")
bullet("Step 3: chat_postMessage — posts the Block Kit summary message with emoji header, date, overall status, and flagged service list")
bullet("Step 4: chat_postMessage — posts the canvas URL as plain text; Slack auto-unfurls slack.com/docs/… as a native Canvas card")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 13. UAA BUSINESS METRICS (uaa_business_collector.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("13. UAA Business Metrics — uaa_business_collector.py")

body(
    "collect_uaa_business_metrics() runs 12 Trino queries concurrently via asyncio.gather. "
    "Each function returns list[BusinessMetric]. Empty returns are silently skipped."
)

table(
    ["Function / SQL File",               "Section",            "metric_type",        "What it measures"],
    [
        ["onboarding_provider_sessions",   "Onboarding",         "provider_comparison","D vs D-1 sessions + success per provider (AKOYA, PLAID, DL_CAPITALONE)"],
        ["account_linking_by_source",      "Account Linking",    "source_comparison",  "Yesterday vs day-before account linkings by client_source × flow_type"],
        ["plaid_batch_recency",            "Plaid Batch Refresh","multi_col_table",     "P50/P75/P90/P95/P99 of hours since last_data_updated_at"],
        ["plaid_batch_metadata_recency",   "Plaid Batch Refresh","multi_col_table",     "Hours since latest run_timestamp in metadata"],
        ["plaid_batch_historical_recency", "Plaid Batch Refresh","multi_col_table",     "P50 data recency trend over last 7 days"],
        ["plaid_batch_trend",              "Plaid Batch Refresh","multi_col_table",     "Daily batch refresh volume trend"],
        ["plaid_batch_refresh_errors",     "Plaid Batch Refresh","multi_col_table",     "Error type breakdown for batch refreshes"],
        ["plaid_force_refresh_daily",      "Plaid Force Refresh","multi_col_table",     "Today vs yesterday force refresh counts"],
        ["plaid_force_refresh_errors",     "Plaid Force Refresh","multi_col_table",     "Error type breakdown for force refreshes"],
        ["plaid_force_refresh_trend",      "Plaid Force Refresh","multi_col_table",     "Multi-day trend of force refresh volume"],
        ["txn_quality_metrics",            "Transaction Quality","multi_col_table",     "Transaction classification quality metrics vs day-before"],
        ["partner_costs",                  "Partner Costs",      "multi_col_table",     "One-time + maintenance + daily total cost per partner, latest available date (MAX run_date)"],
    ],
    col_widths=[2.1, 1.6, 1.5, 2.6],
)

body("SAISM and ALSM latency are fetched from VictoriaMetrics (not Trino):")
bullet("_fetch_alsm_latency(): P50 + P99 + P99 yesterday per aggregator (PLAID, DL_CAPITALONE); computed via histogram_quantile offset 24h")
bullet("_fetch_saism_latency(): same pattern for CRBAA + BRIGHT aggregators; ACCOUNTS_INGESTION_START → ACCOUNTS_CREATED event duration")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 14. DATA PLATFORM — EMR METRICS (emr_collector.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("14. Data Platform EMR Metrics — emr_collector.py")

body(
    "8 SQL queries run concurrently via asyncio.gather. Each produces an EmrSection with a title, "
    "headers, list of EmrRow (cells + flagged bool), and a flag_count. "
    "The EmrReport carries total_flags as a property."
)

table(
    ["Query / SQL File",         "EmrSection Title",                                   "Flag Condition"],
    [
        ["emr_cube_health",       "Cube Health Overview",                               "recency_breach = True"],
        ["emr_staleness",         "Staleness (ordered by staleness_hrs)",               "staleness_hrs > 24"],
        ["emr_memory_top10",      "Memory Usage — Top 10",                              "p95_memory_used_gb > 8.0 GB"],
        ["emr_cpu",               "CPU Utilisation (low utilisation first)",            "p50_cpu_utilization_pct < 10%"],
        ["emr_schedule_delay",    "Schedule Delay (ordered DESC)",                      "p95_schedule_delay_hrs > 1.0h"],
        ["emr_latest_staleness",  "Latest Staleness with Config (ordered DESC)",        "recency_breach OR staleness_hrs > 24"],
        ["emr_row_growth",        "Row Growth (top 50 by new_rows_added)",              "new_rows_added < 0 (row shrinkage)"],
        ["emr_execution_time",    "Execution Time (ordered by P95 DESC)",               "p95_execution_time_hrs > 4.0h"],
    ],
    col_widths=[1.8, 3.0, 3.0],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 15. DATA PLATFORM BUSINESS METRICS (dp_business_collector.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("15. Data Platform Business Metrics — dp_business_collector.py")

body(
    "8 concurrent Trino queries against iceberg_db.cosmos_db__public__* mirror views. "
    "Each returns a BusinessMetric with a details list of affected table/dataset names."
)

table(
    ["SQL File",           "Section",             "What it flags"],
    [
        ["table_recency",      "Table Recency",       "CDC tables with null or stale last-updated timestamps"],
        ["compaction",         "Compaction",          "Iceberg tables that haven't been compacted recently"],
        ["offset_validation",  "Offset Validation",   "Kafka offset validation checks that failed or are stale"],
        ["view_stale",         "View Staleness",       "Cosmos views that haven't refreshed within expected window"],
        ["dbz_invalid",        "Debezium Validation",  "Invalid table events flagged by dp_cosmos_flag_debezium_invalid_tables"],
        ["full_validation",    "Full Validation",      "Full table validation jobs that failed or are overdue"],
        ["base_refresh",       "Base Refresh",         "Base table refresh jobs that are stale or failed"],
        ["validation_stale",   "Validation Staleness", "Validation runs that haven't completed in expected time"],
    ],
    col_widths=[1.6, 1.8, 4.4],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 16. CENTRAL SERVICES BUSINESS METRICS (central_business_collector.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("16. Central Services Business Metrics — central_business_collector.py")

body(
    "Unlike UAA and DP business metrics, Central Services metrics are defined externally "
    "in central_business.json at the project root — not in code. "
    "This allows adding new metrics without a code deploy."
)

body("Each entry in central_business.json has:")
bullet("query_name, display_name, section, metric_type (success_rate | failure_count | total_count | rate)")
bullet("query: PromQL expression")
bullet("taglist (optional): if set, query_vector() is called and each label becomes a separate BusinessMetric")
bullet("warn_below / crit_below: for success_rate — flag if value drops below threshold")
bullet("warn_above / crit_above: for failure_count — flag if value exceeds threshold")

body(
    "All entries run concurrently via asyncio.gather. "
    "Entries returning None from VM are silently skipped — no empty rows in the canvas."
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 17. MODELS (models.py)
# ══════════════════════════════════════════════════════════════════════════════
h1("17. Data Models — models.py")

table(
    ["Dataclass",              "Key Fields"],
    [
        ["Status (enum)",       "HEALTHY, WARNING, CRITICAL, UNKNOWN"],
        ["FlaggingThresholds",  "p99_warn_ms, p99_crit_ms, success_warn_pct, top_n_unflagged"],
        ["ServerMetrics",       "cpu_pct, mem_pct, disk_pct"],
        ["Server",              "name, group, metrics: ServerMetrics, status: Status"],
        ["SystemHealth",        "servers: list[Server]; properties: online, down"],
        ["ApiMetrics",          "throughput_rps, success_rate_pct, error_rate_pct, avg_latency_p50_ms + 3 baselines"],
        ["Endpoint",            "path, hits, success_pct, errors, p99_ms, p99_baseline_ms, success_baseline_pct"],
        ["QueueDepth",          "name, ready, unacked, total"],
        ["ConnectorTask",       "id, state (RUNNING/FAILED/UNASSIGNED/PAUSED)"],
        ["ConnectorStatus",     "name, state, tasks; is_healthy: state==RUNNING and all tasks RUNNING"],
        ["KafkaConnectInstance","name, total, unhealthy: list[ConnectorStatus]"],
        ["KafkaConnectHealth",  "instances: list[KafkaConnectInstance]"],
        ["AirflowDagRun",       "dag_id, state, start_date, end_date, run_date: Optional[date] (IST calendar date)"],
        ["ViewFlowRun",         "table_name, state, start_date"],
        ["ViewFlowHealth",      "total, successful, failed: list[ViewFlowRun], running: list[ViewFlowRun]"],
        ["AirflowHealth",       "dag_runs, view_flow: Optional[ViewFlowHealth], pipeline_runs: list[AirflowDagRun]"],
        ["L0Report",            "service, reported_at, status, system, api, endpoints, queues, thresholds, show_api_metrics"],
    ],
    col_widths=[2.0, 5.8],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 18. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("18. Deployment")

h2("18.1  Build Pipeline")
body(
    "buildspec.yaml defines an AWS CodeBuild pipeline triggered on every git push to main. "
    "It builds the Docker image and pushes to ECR. "
    "No manual image management is required — push = deploy."
)

h2("18.2  Airflow DAG")
table(
    ["Attribute",         "Value"],
    [
        ["DAG ID",         "central_services_agents_l0_report"],
        ["Operator",       "KubernetesPodOperator — spins up a fresh K8s pod from ECR image, runs report, terminates"],
        ["Schedule",       "04:30 UTC daily (10:00 IST)"],
        ["Task ID",        "run_l0_report"],
        ["Image source",   "ECR — rebuilt by CodeBuild on every main branch push"],
        ["Secrets",        "Injected as K8s environment variables: SLACK_BOT_TOKEN, TRINO_USER, VM_AUTH_HEADER, AIRFLOW_DB_URL, etc."],
        ["Logs",           "Stored in S3: eks-airflow-platform/logs/dag_id=central_services_agents_l0_report/..."],
    ],
    col_widths=[1.8, 6.0],
)

h2("18.3  Local Development")
body("Run report immediately:")
code("python -m metrics_report.main --now")
body("Run for one group only:")
code("python -m metrics_report.main --now --group 'UAA Services'")
body("Run on scheduler (blocks):")
code("python -m metrics_report.main")
body("Required: .env file with at minimum SLACK_BOT_TOKEN and SLACK_CHANNEL_ID set.")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# 19. END-TO-END DATA FLOW SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("19. End-to-End Data Flow Summary")

table(
    ["Step", "Component",              "What happens"],
    [
        ["1",  "Airflow DAG",           "KubernetesPodOperator starts the container at 04:30 UTC"],
        ["2",  "main.py",               "asyncio.run(_now()) → run_report()"],
        ["3",  "services.py",           "load_services() parses services.json + ems.json → 30 ServiceDef objects"],
        ["4",  "VMClient.__aenter__",   "httpx.AsyncClient opened with Authorization header from VM_AUTH_HEADER"],
        ["5",  "collector.collect()",   "For each of 30 services: system + API + endpoint + queue + spike PromQL queries through gateway → MetricsReport"],
        ["6",  "formatter.to_l0_report()","MetricsReport → L0Report: baseline comparisons applied, status computed, endpoints structured"],
        ["7",  "spike_analyzer.analyze()","48 × 30-min buckets → SpikeResult (2× jump or 1.5× max/avg)"],
        ["8",  "Trino queries",         "UAA (12) + DP (8) + EMR (8) queries run concurrently via execute_query()"],
        ["9",  "Airflow DB query",      "dag_run table queried for latest pipeline DAG states in IST"],
        ["10", "Kafka Connect API",     "All connector instances polled; unhealthy connectors identified"],
        ["11", "canvas_renderer.render_canvas()","L0Reports + Airflow + Connector health → single markdown string per group"],
        ["12", "slack_publisher.publish_canvas()","canvases.create → canvas_id; chat_postMessage summary; chat_postMessage canvas URL"],
        ["13", "Business canvases",     "Central + UAA + DP business renderers produce and publish their own canvases"],
        ["14", "EMR + DP-L0 canvases",  "emr_renderer + dp_l0_renderer publish their own canvases"],
    ],
    col_widths=[0.4, 2.3, 5.1],
)

divider()

doc.save(OUTPUT)
print(f"✓ Saved: {OUTPUT}")
