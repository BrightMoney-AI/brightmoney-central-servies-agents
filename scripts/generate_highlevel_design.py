"""
generate_highlevel_design.py  —  HL ops channel design spec (revised tier model).
Run: python generate_highlevel_design.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
s = doc.sections[0]
s.top_margin = s.bottom_margin = Cm(2)
s.left_margin = s.right_margin = Cm(2.5)

DARK  = RGBColor(0x1E, 0x29, 0x4C)
BLUE  = RGBColor(0x1F, 0x4E, 0x79)
GREY  = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HDR   = "1F4E79"
ALT   = "EBF5FB"
L0_BG = "EEF0FE"   # indigo tint — trend layer
L1_BG = "EAFAF1"   # green tint  — detail layer
L2_BG = "FEF5E7"   # amber tint  — deep analysis


def _shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = DARK
        r.font.size = Pt(16)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)


def h2(text, color=None):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = color or BLUE
        r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)


def para(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(f"↳  {text}")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)


def tbl(headers, rows, widths=None, tier_col=None):
    """tier_col: column index that holds L0/L1/L2 — tints the row background."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        _shd(c, HDR)
        for para in c.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        if widths:
            c.width = Inches(widths[i])
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        # pick background by tier label
        bg = ALT if ri % 2 == 1 else "FFFFFF"
        if tier_col is not None and ri < len(rows):
            tier = str(rows[ri][tier_col]).strip()
            if tier == "L0":
                bg = L0_BG
            elif tier == "L1":
                bg = L1_BG
            elif tier == "L2":
                bg = L2_BG
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            _shd(c, bg)
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if widths:
                c.width = Inches(widths[ci])
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(50)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("High-Level Ops Channel — Metrics Design")
r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = DARK

for line, sz in [
    ("Canvas Structure & Tier Model", 13),
    ("brightmoney Observability Platform  ·  July 2026", 10),
]:
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(line)
    r2.font.size = Pt(sz); r2.font.color.rgb = GREY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. THE BIG PICTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("1. The Big Picture")

para(
    "Two Slack channels. One existing, one new. The existing channel is untouched — "
    "it carries 8 detailed engineering canvases. The new channel carries 3 canvases "
    "(UAA · Central · Data Platform) designed for leadership and on-call ops: "
    "start at L0 to see if anything is wrong, drop to L1 if it is, "
    "and go to L2 only when you need to understand root cause."
)

tbl(
    ["", "Existing detailed channel", "New high-level channel"],
    [
        ["Canvases / run", "8",                                              "3  (one per team)"],
        ["Audience",       "Service engineers",                               "Leadership · on-call ops"],
        ["What's in it",   "Full endpoint tables, error reasons, EMR cube tables, partner cost rows", "Trend summary → detail → deep analysis, 3 tiers per canvas"],
        ["L0 means",       "Infrastructure (servers)",                        "Trends + instance health — read this first"],
        ["L1 means",       "API health (per endpoint)",                       "Deeper state of what L0 flagged"],
        ["L2 means",       "Business KPIs (full tables)",                     "Root-cause / historical analysis"],
        ["Env var",        "SLACK_CHANNEL_ID",                                "SLACK_HL_CHANNEL_ID  (new, empty = disabled)"],
    ],
    widths=[1.5, 2.5, 2.7],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. TIER MODEL — HOW TO READ A CANVAS
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Tier Model — How to Read a Canvas")

para(
    "The three tiers form a funnel. You are not meant to read all three every day. "
    "L0 tells you whether to keep reading."
)

tbl(
    ["Tier", "Name", "The question it answers", "What's in it", "When to read it"],
    [
        ["L0",
         "Trends & Instance Health",
         "Is something wrong right now?",
         "Overall status per service. Up/down trend on every key metric vs yesterday / 7-day avg. "
         "Server availability. Business flow trends (linking ▲4%, batch success ▼3%). "
         "For DP: CDC table staleness counts, EMR breach count, Debezium invalid table count.",
         "Every day — takes < 1 min to scan"],
        ["L1",
         "Current State Detail",
         "What exactly is wrong?",
         "Per-server CPU / MEM / Disk breakdown. Per-endpoint API breakdown with flagged paths. "
         "Per-sink CDC lag & heartbeat. Connector states. Airflow DAG run status. "
         "Plaid batch hourly trend. Onboarding per provider. Queue depths.",
         "When L0 shows 🟡 or 🔴 — drill into the flagged service"],
        ["L2",
         "Deep Analysis",
         "Why is it wrong? What's the history?",
         "Error reason breakdowns. Historical recency percentiles. Per-partner cost table. "
         "Txn quality by cohort. DP: stale table names list, validation failure details, "
         "EMR cube full table (memory, CPU, execution time). CDC offset validation details.",
         "When L1 doesn't fully explain root cause"],
    ],
    widths=[0.5, 1.8, 2.0, 2.8, 1.6],
    tier_col=0,
)

note(
    "Colour coding in the canvas: L0 sections have an indigo left stripe, "
    "L1 sections green, L2 amber — so you can navigate by eye without reading headings."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. CANVAS STRUCTURE — ALL 3 CANVASES
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Canvas Structure — All 3 Canvases")

para(
    "Every canvas opens with an Attention Required block that lists every flag across "
    "all three tiers. Below that the tiers appear in order: L0 → L1 → L2. "
    "A reader who only needs the headline can stop after L0."
)

tbl(
    ["#", "Section", "Tier", "Content", "Skipped when"],
    [
        ["1",  "Canvas title",          "—",  "<Group> — Health Overview — DD Mon YYYY",                             "Never"],
        ["2",  "⚠️ Attention Required", "—",  "Every flag from L0 + L1 + L2, critical first then warning",          "No flags across any tier"],
        ["3",  "L0  Trends & Health",   "L0", "Per-service status + trend lines for every key metric",               "Never (always rendered)"],
        ["4",  "L1  Detail",            "L1", "Per-server breakdown · per-endpoint breakdown · queues · CDC sinks · connectors · Airflow", "No L1 data available"],
        ["5",  "L2  Deep Analysis",     "L2", "Error reasons · historical tables · root-cause data",                  "No L2 data returned"],
        ["6",  "Legend",                "—",  "🟢 Healthy  🟡 Warning  🔴 Critical  ·  brightmoney observability",  "Never"],
    ],
    widths=[0.3, 1.8, 0.5, 3.7, 1.4],
)

note(
    "L1 and L2 sections are collapsible in Slack Canvas. "
    "On a healthy day a reader sees L0, scans all green, and stops. "
    "L1 and L2 are there when needed — not noise on healthy days."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. FLAGGING RULES — ATTENTION REQUIRED
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Flagging Rules — Attention Required Block")

para(
    "Any metric that breaches a threshold generates one line in Attention Required. "
    "Format:  🟡/🔴  Service  ·  Tier  ·  Metric  ·  Value  (context e.g. '▲ +42% vs 7d'). "
    "Lines sorted: 🔴 critical first, then 🟡 warning; within each severity sorted by service."
)

h2("L0 flag conditions")
tbl(
    ["Metric", "🟡 Warning", "🔴 Critical", "Trend shown"],
    [
        ["API success rate",    "Drop ≥ 5 pp vs 7d avg",    "Drop ≥ 10 pp vs 7d avg",    "▲/▼ N pp vs 7d"],
        ["API error rate",      "Spike ≥ 2× vs 7d avg",     "Spike ≥ 3× vs 7d avg",      "▲ N× vs 7d"],
        ["API latency P50",     "≥ 1.5× 7d baseline",       "≥ 2.0× 7d baseline",        "▲/▼ +N% vs 7d"],
        ["Linking success",     "—  (trend only)",           "—  (trend only)",            "▲/▼ N% vs D-1"],
        ["Onboarding sessions", "—  (trend only)",           "—  (trend only)",            "▲/▼ N% vs D-1"],
        ["Plaid batch success", "< 95%",                     "< 90%",                      "▼ N pp vs last hour"],
        ["ALSM / SAISM P99",    "Increased vs yesterday",    "—",                          "▲ +Ns vs yesterday"],
        ["Server availability", "—",                         "Any node down / no data",    "N online / M down"],
        ["DP: stale CDC tables","—",                         "> 0 stale tables",           "N stale"],
        ["DP: EMR cube breach", "—",                         "> 0 cubes breached",         "N breached"],
        ["DP: DBZ invalid tables","—",                       "> 0 invalid tables",         "N invalid"],
        ["DP: CDC lag trend",   "Growing > 100 msgs / 24h",  "Growing > 5,000 msgs / 24h", "▲ +N msgs / 24h"],
    ],
    widths=[1.7, 1.6, 1.6, 2.3],
)

h2("L1 flag conditions")
tbl(
    ["Metric", "🟡 Warning", "🔴 Critical"],
    [
        ["CPU % per server",          "> 70%",                    "> 90%"],
        ["Memory % per server",       "> 75%",                    "> 90%"],
        ["Disk % per server",         "> 80%",                    "> 90%"],
        ["Endpoint success rate",     "Drop ≥ 5 pp vs 7d",        "Drop ≥ 10 pp vs 7d  (≥ 100 hits)"],
        ["Endpoint P99 latency",      "≥ 1.5× 7d baseline",       "≥ 2.0× 7d baseline  (≥ 100 hits)"],
        ["RabbitMQ ready msgs",       "≥ 100 msgs ready",         "≥ 500 msgs ready"],
        ["CDC coord lag per sink",    "> 1,000 msgs",              "> 10,000 msgs"],
        ["CDC heartbeat",             "—",                         "< 5 msgs in 5-min window"],
        ["Kafka Connect state",       "—",                         "Any connector not RUNNING"],
        ["Airflow DAG run",           "—",                         "state = failed"],
        ["Airflow view-flow refresh", "—",                         "Any failed table refresh in 24 h"],
        ["Iceberg / Debezium disk %", "> 80%",                    "> 90%"],
    ],
    widths=[2.5, 2.2, 3.0],
)

h2("L2 flag conditions")
tbl(
    ["Metric", "Flag condition"],
    [
        ["Plaid batch error reasons",   "Any error class with count > 0 in last 24 h"],
        ["DP compaction needed",        "Any table with file growth > 300 in 3 days"],
        ["DP offset / full validation", "Any validation failure count > 0"],
        ["DP base refresh failures",    "Any failure count > 0"],
        ["DP view staleness",           "Any stale view count > 0"],
        ["EMR cube staleness",          "Any cube with staleness_hrs > 24"],
        ["EMR schedule delay",          "Any cube with P95 schedule delay > 1 h"],
    ],
    widths=[2.5, 5.2],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. PER-CANVAS BREAKDOWN
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Per-Canvas Metric Breakdown")

para(
    "Each table below lists every metric in that canvas with its tier, "
    "what it shows, and whether it generates a flag. "
    "Rows are colour-coded: L0 = indigo, L1 = green, L2 = amber."
)

# ── Canvas 1: UAA ─────────────────────────────────────────────────────────────
h2("Canvas 1 — UAA Services")

tbl(
    ["Tier", "Metric", "What it shows", "Flag?"],
    [
        # L0
        ["L0", "Service overall status",          "🟢/🟡/🔴 per service — worst signal across all metrics",                   "Always shown"],
        ["L0", "Server availability",              "N online / M down per service",                                            "🔴 any node down"],
        ["L0", "API success rate + trend",         "Success % with ▲/▼ vs 7d avg",                                            "🟡 −5 pp  🔴 −10 pp vs 7d"],
        ["L0", "API error rate + trend",           "Error % with ▲/▼ vs 7d avg",                                              "🟡 2×  🔴 3× vs 7d"],
        ["L0", "API latency P50 + trend",          "Median latency with ▲/▼ vs 7d avg",                                       "🟡 1.5×  🔴 2× vs 7d"],
        ["L0", "Onboarding sessions trend",        "Total success sessions today vs D-1  (▲/▼ %)",                            "Trend only — informational"],
        ["L0", "Account linking trend",            "Successful linkings D-1 vs D-2  (▲/▼ count + %)",                         "Trend only — informational"],
        ["L0", "ALSM latency P99 trend",           "Worst P99 PLAID / DL_CAPITALONE vs yesterday  (▲/▼ Ns)",                  "🟡 increased vs yesterday"],
        ["L0", "SAISM latency P99 trend",          "Worst P99 CRBAA / BRIGHT vs yesterday  (▲/▼ Ns)",                         "🟡 increased vs yesterday"],
        ["L0", "Plaid batch success trend",        "Success % latest hour  (▲/▼ vs prior hour)",                              "🟡 <95%  🔴 <90%"],
        ["L0", "Plaid force refresh success trend","Success % yesterday  (▲/▼ vs D-2)",                                       "🟡 <95%  🔴 <90%"],
        # L1
        ["L1", "CPU / Memory / Disk per server",  "Per-node breakdown — group summary + individual rows",                     "🟡/🔴 per threshold"],
        ["L1", "Flagged endpoints",               "Endpoints with success drop or latency spike vs 7d  (≥ 100 hits)",          "🟡 −5 pp / 1.5×   🔴 −10 pp / 2×"],
        ["L1", "Top 5 endpoints by volume",       "Healthy high-traffic paths — hits, success %, P99",                        "Informational"],
        ["L1", "Queue depth (RabbitMQ)",          "Ready + unacked per queue name",                                           "🟡 ready ≥100  🔴 ready ≥500"],
        ["L1", "Onboarding by provider",          "AKOYA / PLAID / DL_CAPITALONE — sessions, success count, success % today vs D-1", "Informational"],
        ["L1", "Account linking by source & flow","Sessions by client_source × flow_type: D-1 vs D-2 + delta",               "Informational"],
        ["L1", "ALSM / SAISM latency table",      "P50 + P99 today and yesterday per aggregator",                            "Flagged if P99 increased"],
        ["L1", "Plaid batch hourly trend",        "Success % and error % per hour — last 24 h",                              "Flagged if success < 95%"],
        # L2
        ["L2", "Plaid batch error reasons",       "Top error classes by count — last 24 h, grouped by hour",                  "🔴 any errors present"],
        ["L2", "Plaid batch historical recency",  "P50/P75/P90/P95/P99 recency percentiles — last 2 days",                   "Informational"],
        ["L2", "Plaid force refresh error reasons","Top error classes — last 7 days",                                         "🔴 any errors present"],
        ["L2", "Partner cost breakdown",          "Per partner: one-time cost, maintenance cost, daily total  (yesterday)",   "Informational"],
        ["L2", "Txn quality by cohort",           "Avg and P95 txn duration (days) and txn count per account cohort × provider","Informational"],
    ],
    widths=[0.5, 2.3, 2.9, 2.0],
    tier_col=0,
)

doc.add_page_break()

# ── Canvas 2: Central ─────────────────────────────────────────────────────────
h2("Canvas 2 — Central Services")

tbl(
    ["Tier", "Metric", "What it shows", "Flag?"],
    [
        # L0
        ["L0", "Service overall status",          "🟢/🟡/🔴 per service",                                                  "Always shown"],
        ["L0", "Server availability",              "N online / M down per service",                                          "🔴 any node down"],
        ["L0", "API success rate + trend",         "Success % with ▲/▼ vs 7d",                                              "🟡 −5 pp  🔴 −10 pp"],
        ["L0", "API error rate + trend",           "Error % with ▲/▼ vs 7d",                                                "🟡 2×  🔴 3×"],
        ["L0", "API latency P50 + trend",          "Latency with ▲/▼ vs 7d",                                                "🟡 1.5×  🔴 2×"],
        ["L0", "Business metrics scorecard",       "Per section from central_business.json: N flagged / total checks",       "🟡/🔴 per metric's own threshold"],
        # L1
        ["L1", "CPU / Memory / Disk per server",  "Per-node breakdown for all Central nodes",                               "🟡/🔴 per threshold"],
        ["L1", "Flagged endpoints",               "Endpoints with success drop or latency spike vs 7d  (≥ 100 hits)",        "🟡/🔴 per threshold"],
        ["L1", "Top 5 endpoints by volume",       "Healthy high-traffic paths",                                              "Informational"],
        ["L1", "Central business metric values",  "Per-section current values — every metric from central_business.json",   "🟡/🔴 per metric threshold"],
        # L2
        ["L2", "Full central business tables",    "Historical context for any flagged business metric sections",             "If section had flags in L1"],
    ],
    widths=[0.5, 2.3, 2.9, 2.0],
    tier_col=0,
)

note(
    "Forwarder services (Mixpanel, Singular, Facebook, Snap, Google) have no api_job. "
    "They appear in L0 server availability only — no API trend rows."
)

doc.add_page_break()

# ── Canvas 3: Data Platform ───────────────────────────────────────────────────
h2("Canvas 3 — Data Platform")

tbl(
    ["Tier", "Metric", "What it shows", "Flag?"],
    [
        # L0
        ["L0", "Service overall status",            "🟢/🟡/🔴 per DP service",                                            "Always shown"],
        ["L0", "Server availability",               "N online / M down per service",                                        "🔴 any node down"],
        ["L0", "API success rate + trend",          "Success % with ▲/▼ vs 7d  (DP services with api_job)",                 "🟡 −5 pp  🔴 −10 pp"],
        ["L0", "API error rate + trend",            "Error % with ▲/▼ vs 7d",                                              "🟡 2×  🔴 3×"],
        ["L0", "API latency P50 + trend",           "Latency with ▲/▼ vs 7d",                                              "🟡 1.5×  🔴 2×"],
        ["L0", "Stale / null CDC table count",      "How many CDC tables have null or stale recency right now",              "🔴 any stale table"],
        ["L0", "Debezium invalid table count",      "How many Debezium tables are currently in an invalid state",            "🔴 any invalid table"],
        ["L0", "EMR cube recency breach count",     "How many EMR cubes have missed their recency SLA",                     "🔴 any breach"],
        ["L0", "Compaction needed count",           "Tables with file growth > 300 in last 3 days",                         "🔴 any table"],
        ["L0", "CDC lag trend (aggregate)",         "Is total CDC offset lag growing or stable across all sinks?",           "🟡 growing >100  🔴 >5k msgs/24h"],
        ["L0", "Airflow DAG health summary",        "Overall DAG success count vs failed count — today vs yesterday",        "🔴 any DAG failed"],
        # L1
        ["L1", "CPU / Memory / Disk per server",   "Per-node breakdown for all DP service nodes",                           "🟡/🔴 per threshold"],
        ["L1", "Iceberg / Debezium VM disk %",      "Disk usage on CDC infrastructure VMs",                                  "🟡 >80%  🔴 >90%"],
        ["L1", "Flagged endpoints",                "Endpoints with success drop or latency spike  (≥ 100 hits)",             "🟡/🔴 per threshold"],
        ["L1", "CDC per-sink detail",              "Per sink: coord lag, sum offset lag, lag delta 24h, heartbeat msgs/5m", "🟡/🔴 per sink threshold"],
        ["L1", "Kafka Connect per-connector",      "State per connector + task state  (one row per unhealthy connector)",    "🔴 any not RUNNING"],
        ["L1", "Airflow DAG runs today / yesterday","Per-DAG state + start time for today and yesterday",                   "🔴 state = failed"],
        ["L1", "View-flow 24 h summary",           "Total runs · success · failed · running — with failed table list",      "🔴 any failed refresh"],
        ["L1", "Stale CDC table names",            "List of which specific tables are stale or have null recency",           "Shown when L0 count > 0"],
        ["L1", "Debezium invalid table names",     "List of which specific tables are in invalid state",                     "Shown when L0 count > 0"],
        ["L1", "EMR cube breach names",            "Which specific cubes are breaching recency SLA",                        "Shown when L0 count > 0"],
        # L2
        ["L2", "Full CDC offset validation detail","Per-table validation failures with context",                             "🔴 any failure"],
        ["L2", "Full base refresh failure detail", "Which tables failed base refresh and when",                              "🔴 any failure"],
        ["L2", "Full view staleness detail",       "Stale view names, last refresh time",                                   "🔴 any stale view"],
        ["L2", "EMR cube full health table",       "All cubes: last run, age (h), recency, breach, total rows",             "Rendered always for DP"],
        ["L2", "EMR staleness detail",             "Cubes ordered by staleness_hrs with last run time",                     "Rendered always for DP"],
        ["L2", "EMR memory top 10",                "Top 10 cubes by P95 memory usage (GB)",                                 "Rendered always for DP"],
        ["L2", "EMR execution time",               "P50/P95 execution time per cube, ordered by P95 desc",                  "Rendered always for DP"],
        ["L2", "EMR schedule delay",               "P50/P95 schedule delay per cube — identifies chronically late cubes",   "🔴 P95 delay > 1 h"],
    ],
    widths=[0.5, 2.3, 2.7, 2.2],
    tier_col=0,
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. IMPLEMENTATION SCOPE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Implementation Scope")

para(
    "The existing 8-canvas flow is completely unchanged. "
    "New code is additive only — two new files plus minimal config changes."
)

h2("New files")
tbl(
    ["File", "Purpose"],
    [
        ["metrics_report/hl_canvas_renderer.py",
         "render_hl_canvas(reports, biz_metrics, dp_extras) → canvas markdown. "
         "Builds Attention Required (all tiers) → L0 trend section → L1 detail section → L2 deep analysis section."],
        ["metrics_report/hl_scheduler.py",
         "run_hl_report() — same collect() loop as run_report(), "
         "calls hl_canvas_renderer, posts to SLACK_HL_CHANNEL_ID. "
         "Reuses all existing collectors unchanged."],
    ],
    widths=[2.8, 4.9],
)

h2("Files to modify  (two lines of change each)")
tbl(
    ["File", "Change"],
    [
        ["metrics_report/config.py", "Add  slack_hl_channel_id: str = ''"],
        ["metrics_report/main.py",   "Start hl_scheduler if SLACK_HL_CHANNEL_ID is set — runs alongside existing scheduler"],
    ],
    widths=[2.3, 5.4],
)

h2("Files that must not change")
tbl(
    ["File / group", "Why it stays unchanged"],
    [
        ["scheduler.py  +  canvas_renderer.py",         "Existing 8-canvas detailed flow — zero risk"],
        ["collector.py  ·  formatter.py  ·  models.py", "Data collection and L0Report builder — shared, read-only"],
        ["vm_client.py  ·  gateway.py  ·  services.py", "Infrastructure — shared"],
        ["slack_publisher.py",                           "Canvas publisher — reused as-is, channel ID passed as param"],
        ["All business / EMR / DP-L0 collectors",        "Same collectors feed both channels — no duplication"],
        ["services.json  /  ems.json",                   "Service definitions — no change"],
    ],
    widths=[3.0, 4.7],
)

h2("One new env var")
tbl(
    ["Env var", "Default", "Description"],
    [
        ["SLACK_HL_CHANNEL_ID", "''  (empty)",
         "Slack channel ID for the high-level ops canvas. "
         "Leave empty to keep HL flow disabled. Existing channel is unaffected."],
    ],
    widths=[1.8, 1.0, 4.9],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. DATA FLOW
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Data Flow")

para(
    "Both channels share the same collection pipeline. "
    "The split happens only at the render + publish step."
)

tbl(
    ["Step", "What happens", "Output"],
    [
        ["1  Schedule",          "APScheduler CronTrigger fires at 04:30 UTC (10:00 IST)",                                       "Calls run_report() + run_hl_report()"],
        ["2  Load services",     "load_services() reads services.json + ems.json",                                                "list[ServiceDef]"],
        ["3  Collect L0 / L1",  "Single VMClient session — queries VM for each service via MetricsGateway (5s timeout per query)","MetricsReport per service"],
        ["4  Format",            "to_l0_report() converts raw MetricsReport → typed L0Report with Status, servers, endpoints",    "L0Report per service"],
        ["5  DP extras in VM",  "dp_l0_collector: CDC lag, heartbeat, sink disk — inside the same VM session",                    "DPL0Report"],
        ["6  Collect L2",       "Trino queries: UAA biz, Central biz, DP biz, EMR — run concurrently outside VM session",        "BusinessMetric lists, EmrReport"],
        ["7  Collect connectors + Airflow", "asyncio.gather: Kafka Connect REST + Airflow REST — concurrent",                    "KafkaConnectHealth, AirflowHealth"],
        ["8a  Render detailed",  "canvas_renderer.render_canvas() per group → publish_canvas() to SLACK_CHANNEL_ID",              "8 canvases → detailed channel"],
        ["8b  Render HL",        "hl_canvas_renderer.render_hl_canvas() per group (L0 trends → L1 detail → L2 analysis) → publish_canvas() to SLACK_HL_CHANNEL_ID", "3 canvases → HL channel"],
    ],
    widths=[1.8, 4.2, 1.8],
)

note(
    "Steps 1–7 run once and feed both channels. "
    "Steps 8a and 8b are the only point where the two channels diverge — "
    "same data, different rendering."
)


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save("HL_Channel_Metrics_Design.docx")
print("✅  Saved: HL_Channel_Metrics_Design.docx")
